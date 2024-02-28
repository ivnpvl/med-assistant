from docx import Document
from odf import text, teletype
from odf.opendocument import load
from functools import cached_property
from pathlib import Path
from typing import Callable

import templates
from config import ARCHIVE_DIR, CARD_DIR, JSON_DIR
from exceptions import StringInvalidError, StringNotExistsError
from normalizer import normalize_date, normalize_name


class File:

    def __init__(self, path: Path):
        self.path = path
        self.suffix = path.suffix
        if self.suffix not in (".docx", ".odt"):
            raise NotImplementedError(
                "Поддерживаются только файлы с расширением .docx или .odt."
            )
        self.data = {"path": str(path)}

    @cached_property
    def document(self):
        if self.suffix == ".docx":
            return Document(self.path)
        if self.suffix == ".odt":
            return load(self.path)

    @property
    def paragraphs(self):
        if self.suffix == ".docx":
            return self.document.paragraphs
        if self.suffix == ".odt":
            return self.document.getElementsByType(text.P)

    def get_text(self, paragraph) -> str:
        if self.suffix == ".docx":
            return paragraph.text
        if self.suffix == ".odt":
            return teletype.extractText(paragraph)

    @staticmethod
    def strip_bolders(text: str, left: str, right: str) -> str:
        return text.split(left)[1].split(right)[0].strip()

    def parse_startwith(self, startwith: str, endwith="\n") -> str:
        for paragraph in self.paragraphs:
            text = self.get_text(paragraph)
            if startwith in text:
                parsed = self.strip_bolders(text, startwith, endwith)
                if not parsed:
                    raise StringInvalidError(startwith)
                return parsed
        raise StringNotExistsError(startwith)

    def edit_startwith(self, startwith: str, edit_func: Callable) -> bool:
        if self.suffix == ".docx":
            for paragraph in self.paragraphs:
                text = paragraph.text
                if startwith in text:
                    parsed = self.strip_bolders(text, startwith, endwith="\n")
                    if not parsed:
                        raise StringInvalidError(startwith)
                    edited = edit_func(parsed)
                    if parsed != edited:
                        paragraph.text = paragraph.text.replace(parsed, edited)
                        return True
                    return False
            raise StringNotExistsError(startwith)
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")

    def add_data(self, attr: str, template: str) -> None:
        if attr == "address":
            parsed = self.parse_startwith(template, endwith=", тел.:")
        parsed = self.parse_startwith(template)
        if attr == "name":
            surname, name, patronymic = parsed.split()
            self.data["surname"] = surname
            self.data["name"] = name
            self.data["patronymic"] = patronymic
        else:
            self.data[attr] = parsed

    def save(self):
        if self.suffix == ".docx":
            self.document.save(self.path)
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")


class Card(File):

    WORK_DIR = CARD_DIR
    JSON_PATH = JSON_DIR / "cards.json"
    FILENAME_ATTRS = templates.CARD_FILENAME_ATTRS
    FILENAME_TEMPLATE = templates.CARD_FILENAME_TEMPLATE
    PARSE_RANGE = templates.CARD_PARSE_RANGE
    PATH_SIGNS = templates.CARD_PATH_SIGNS


class Consultation(File):

    WORK_DIR = ARCHIVE_DIR
    JSON_PATH = JSON_DIR / "consultations.json"
    NORMALIZE_FUNCS = {
        "name": normalize_name,
        "birthdate": lambda data: normalize_date(data).strftime("%d.%m.%Y"),
    }
    PARSE_RANGE = templates.CONSULTATION_PARSE_RANGE
