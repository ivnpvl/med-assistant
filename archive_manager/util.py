from docx import Document
from odf import text, teletype
from odf.opendocument import load
from pathlib import Path
from typing import Callable

from config import ARCHIVE_DIR, CARD_DIR, JSON_DIR
from exceptions import AttrNotExistsError
from normalizer import normalize_date, normalize_name
from templates import (
    CARD_PATH_TEMPLATES,
    CARD_STARTWITH_TEMPLATES,
    CARD_TITLE_ATTRS,
    CARD_TITLE_TEMPLATE,
    CONSULTATION_STARTWITH_TEMPLATES,
)


class File:

    def __init__(self, path: Path):
        self.path = path
        self.suffix = self.path.suffix
        self.document = self._get_document()
        self.paragraphs = self._get_paragraphs()
        self.data = {"path": str(path)}

    def _get_document(self):
        if self.suffix == ".docx":
            return Document(self.path)
        if self.suffix == ".odt":
            return load(self.path)

    def _get_paragraphs(self):
        if self.suffix == ".docx":
            return self.document.paragraphs
        if self.suffix == ".odt":
            return self.document.getElementsByType(text.P)

    def get_text(self, paragraph) -> str:
        if self.suffix == ".docx":
            return paragraph.text
        if self.suffix == ".odt":
            return teletype.extractText(paragraph)

    def parse_startwith(self, template: str) -> str | None:
        for paragraph in self.paragraphs:
            text = self.get_text(paragraph)
            if template in text:
                return text.split(template)[1].split("\n")[0].strip()

    def edit_startwith(self, template: str, normalize_func: Callable) -> bool:
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")
        for paragraph in self.paragraphs:
            text = self.get_text(paragraph)
            if template in text:
                parsed = text.split(template)[1].split("\n")[0].strip()
                if not parsed:
                    raise
                normalized = normalize_func(parsed)
                if parsed != normalized:
                    paragraph.text = paragraph.text.replace(parsed, normalized)
                    return True
        return False

    def add_data(self, attr: str, template: str) -> None:
        parsed = self.parse_startwith(template)
        if not parsed:
            raise AttrNotExistsError(attr)
        if attr == "name":
            surname, name, patronymic = parsed.split()
            self.data["surname"] = surname
            self.data["name"] = name
            self.data["patronymic"] = patronymic
        else:
            self.data[attr] = parsed

    def save(self):
        if self.suffix == ".docx":
            self.document.save(self.path)
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")


class Card(File):

    WORK_DIR = CARD_DIR
    JSON_PATH = JSON_DIR / "cards.json"
    STARTWITH_TEMPLATES = CARD_STARTWITH_TEMPLATES
    PATH_TEMPLATES = CARD_PATH_TEMPLATES
    TITLE_TEMPLATE = CARD_TITLE_TEMPLATE
    TITLE_ATTRS = CARD_TITLE_ATTRS


class Consultation(File):

    WORK_DIR = ARCHIVE_DIR
    JSON_PATH = JSON_DIR / "consultations.json"
    STARTWITH_TEMPLATES = CONSULTATION_STARTWITH_TEMPLATES
    NORMALIZE_FUNCS = {
        "name": normalize_name,
        "birthdate": lambda data: normalize_date(data).strftime("%d.%m.%Y"),
    }


class PercentageScale:

    def __init__(self, number: int, percent: int = 5):
        self.number = number
        self.percent = percent
        self.scale = self._get_percentage_scale()

    def _get_percentage_scale(self) -> dict:
        steps = [step for step in range(self.percent, 100, self.percent)]
        scale = {self.number * step // 100: step for step in steps}
        return scale

    def display(self, number):
        if number in self.scale:
            print(f"Обработано {self.scale[number]}% файлов.")
