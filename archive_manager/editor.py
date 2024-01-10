from docx import Document
from pathlib import Path

from logs.decorator import log_it
from normalizers.date import normalize_date
from normalizers.name import normalize_name
from settings import ARCHIVE_DIR


def get_percentage_scale(number, percent):
    steps = [step for step in range(percent, 100, percent)]
    return {number * step // 100: step for step in steps}


def edit_name_docx(document):
    name_template = "Пациент:   "
    for paragraph in document.paragraphs:
        if name_template in paragraph.text:
            name = paragraph.text.split(name_template)[1].split("\n")[0]
            normalized_name = normalize_name(name)
            if name != normalized_name:
                paragraph.text = paragraph.text.replace(name, normalized_name)
                return True
            return False
    return False


def edit_date_docx(document):
    date_template = "Дата рождения:   "
    for paragraph in document.paragraphs:
        if date_template in paragraph.text:
            date = paragraph.text.split(date_template)[1].split("\n")[0]
            normalized_date = normalize_date(date).strftime("%d.%m.%Y")
            if date != normalized_date:
                paragraph.text = paragraph.text.replace(date, normalized_date)
                return True
            return False
    return False


@log_it
def edit_and_save(path):
    document = Document(path)
    name_changed = edit_name_docx(document)
    date_changed = edit_date_docx(document)
    if name_changed or date_changed:
        document.save(path)


def main():
    paths = list(map(str, Path(ARCHIVE_DIR).glob('*.docx')))
    status_bar = get_percentage_scale(len(paths), 10)
    for counter, path in enumerate(paths, 1):
        edit_and_save(path)
        if counter in status_bar:
            print(f"Отредактировано {status_bar[counter]}% файлов.")


if __name__ == "__main__":
    main()
