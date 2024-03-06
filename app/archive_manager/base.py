from docx import Document
from odf import text, teletype
from odf.opendocument import load
from pathlib import Path

from archive_manager import templates
from archive_manager.exceptions import StringInvalidError, StringNotExistsError
from archive_manager.normalizer import normalize_date, normalize_name
from core.config import ARCHIVE_DIR, CARD_DIR, JSON_DIR


class ArchiveFile:

    def __init__(self, path: Path) -> None:
        self.path = path
        self.suffix = path.suffix
        if self.suffix not in (".docx", ".odt"):
            raise NotImplementedError(
                "Поддерживаются только файлы с расширением .docx или .odt."
            )
        self.document = self._load_document()
        self.paragraphs = self._get_paragraphs()

    def _load_document(self):
        if self.suffix == ".docx":
            return Document(self.path)
        if self.suffix == ".odt":
            return load(self.path)

    def _get_paragraphs(self):
        if self.suffix == ".docx":
            return self.document.paragraphs
        if self.suffix == ".odt":
            return self.document.getElementsByType(text.P)

    def extract_text(self, paragraph) -> str:
        if self.suffix == ".docx":
            return paragraph.text
        if self.suffix == ".odt":
            return teletype.extractText(paragraph)

    @staticmethod
    def strip_bolders(text: str, startwith: str, endwith: str = "\n") -> str:
        return text.split(startwith)[1].split(endwith)[0].strip()

    def parse(self, frame: tuple[str]) -> str:
        for paragraph in self.paragraphs:
            text = self.extract_text(paragraph)
            if frame[0] in text:
                parsed = self.strip_bolders(text, *frame)
                if not parsed:
                    raise StringInvalidError(frame)
                return parsed
        raise StringNotExistsError(frame)

    def extract_data(self) -> dict:
        if not hasattr(self, "PARSE_FRAMES"):
            raise NotImplementedError(
                "Необходимо задать PARSE_FRAMES в дочернем классе."
            )
        data = {}
        for attr, frame in self.PARSE_FRAMES.items():
            parsed = self.parse(frame)
            if attr == "fullname":
                try:
                    surname, name, patronymic = parsed.split()
                except ValueError:
                    raise StringInvalidError(frame)
                data["surname"] = surname
                data["name"] = name
                data["patronymic"] = patronymic
            else:
                data[attr] = parsed
        return data

    def edit_startwith(self, startwith: str, edit_func) -> bool:
        if self.suffix == ".docx":
            for paragraph in self.paragraphs:
                text = paragraph.text
                if startwith in text:
                    parsed = self.strip_bolders(text, startwith, endwith="\n")
                    if not parsed:
                        raise StringInvalidError(startwith)
                    edited = edit_func(parsed)
                    if parsed != edited:
                        paragraph.text = paragraph.text.replace(parsed, edited)
                        return True
                    return False
            raise StringNotExistsError(startwith)
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")

    def save(self):
        if self.suffix == ".docx":
            self.document.save(self.path)
        if self.suffix == ".odt":
            raise NotImplementedError("Файл .odt не поддерживает изменение.")


class Card(ArchiveFile):

    WORK_DIR = CARD_DIR
    JSON_PATH = JSON_DIR / "cards.json"
    FILENAME_ATTRS = templates.CARD_FILENAME_ATTRS
    FILENAME_TEMPLATE = templates.CARD_FILENAME_TEMPLATE
    PARSE_FRAMES = templates.CARD_PARSE_FRAMES
    PATH_SIGNS = templates.CARD_PATH_SIGNS


class Consultation(ArchiveFile):

    WORK_DIR = ARCHIVE_DIR
    JSON_PATH = JSON_DIR / "consultations.json"
    NORMALIZE_FUNCS = {
        "name": normalize_name,
        "birthdate": lambda data: normalize_date(data).strftime("%d.%m.%Y"),
    }
    PARSE_FRAMES = templates.CONSULTATION_PARSE_FRAMES
