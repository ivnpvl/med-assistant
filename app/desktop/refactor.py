import os
import app.desktop.checkbuttons as text
import tkinter as tk
from tkinter import messagebox as mb
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime



def ask_about_card():
    answer = mb.askyesno(title="Вопрос", message="Завести амбулаторную карту?")
    if answer:
        root.destroy()
        card_creation()
    else:
        root.destroy()


# create card
def card_creation():
    global root, ent_sex, ent_region, ent_district, ent_city, ent_street, ent_tel
    root = tk.Tk()
    root.title("Амбулаторная карта")
    root.geometry("1280x640+40+40")
    root.resizable(False, False)
    maincolor = "#A8DE56"
    root.config(bg=maincolor)

    lbl_sex = tk.Label(root, text="Пол:", bg=maincolor, font=("Times New Roman", 12))
    ent_sex = tk.Entry(root, font=("Times New Roman", 12), width=8)
    lbl_region = tk.Label(root, text="Область:", bg=maincolor, font=("Times New Roman", 12))
    ent_region = tk.Entry(root, font=("Times New Roman", 12), width=50)
    ent_region.insert(0, "Пензенская обл.")
    lbl_district = tk.Label(root, text="Район:", bg=maincolor, font=("Times New Roman", 12))
    ent_district = tk.Entry(root, font=("Times New Roman", 12), width=50)
    lbl_city = tk.Label(root, text="Населённый пункт:", bg=maincolor, font=("Times New Roman", 12))
    ent_city = tk.Entry(root, font=("Times New Roman", 12), width=50)
    ent_city.insert(0, "г. Пенза")
    lbl_street = tk.Label(root, text="Адрес (ул. Улица, дом-кв.):", bg=maincolor, font=("Times New Roman", 12))
    ent_street = tk.Entry(root, font=("Times New Roman", 12), width=50)
    lbl_tel = tk.Label(root, text="Телефон:", bg=maincolor, font=("Times New Roman", 12))
    ent_tel = tk.Entry(root, font=("Times New Roman", 12), width=50)
    btn = tk.Button(root, text="Готово", bd=5, width=10, height=2, font=("Times New Roman", 12), command=get_entry_card)

    lbl_sex.grid(row=0, column=0, padx=80, pady=5, sticky="w")
    ent_sex.grid(row=0, column=1, pady=5, sticky="w")
    lbl_region.grid(row=1, column=0, padx=80, pady=5, sticky="w")
    ent_region.grid(row=1, column=1, pady=5, sticky="w")
    lbl_district.grid(row=2, column=0, padx=80, pady=5, sticky="w")
    ent_district.grid(row=2, column=1, pady=5, sticky="w")
    lbl_city.grid(row=3, column=0, padx=80, pady=5, sticky="w")
    ent_city.grid(row=3, column=1, pady=5, sticky="w")
    lbl_street.grid(row=4, column=0, padx=80, pady=5, sticky="w")
    ent_street.grid(row=4, column=1, pady=5, sticky="w")
    lbl_tel.grid(row=5, column=0, padx=80, pady=5, sticky="w")
    ent_tel.grid(row=5, column=1, pady=5, sticky="w")
    ent_tel.grid(row=5, column=1, pady=5, sticky="w")
    btn.grid(row=6, column=0, padx=80, pady=5, sticky="w")


def get_entry_card():
    global sex, registration, add_card
    add_card = True
    sex = ent_sex.get().strip()
    if sex == "м" or sex == "м." or sex == "муж":
        sex = "муж."
    if sex == "ж" or sex == "ж." or sex == "жен":
        sex = "жен."
    region = ent_region.get().strip()
    if region != "":
        region = region + ", "
    district = ent_district.get().strip()
    if district != "":
        district = district + ", "
    city = ent_city.get().strip()
    if city != "":
        city = city + ", "
    street = ent_street.get().strip()
    if street != "":
        street = street + ", "
    tel = ent_tel.get().replace(" ", "")
    if len(tel) == 6:
        tel = f"тел.: (+7 8412) {tel[:2]}-{tel[2:4]}-{tel[4:6]}"
    elif len(tel) == 10:
        tel = f"тел.: +7 {tel[:3]} {tel[3:6]}-{tel[6:8]}-{tel[8:10]}"
    elif len(tel) == 11:
        tel = f"тел.: +7 {tel[1:4]} {tel[4:7]}-{tel[7:9]}-{tel[9:11]}"
    else:
        tel = f"тел.: {tel}"
    registration = region + district + city + street + tel
    root.destroy()


# page config
document = Document()
section = document.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section = document.sections[-1]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
paragraph_format = document.styles["Normal"].paragraph_format
paragraph_format.line_spacing = Pt(13)
style = document.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(10)

# date in dd.mm.yyyy
current_datetime = datetime.now()
if current_datetime.day > 9:
    day = f"{current_datetime.day}"
else:
    day = f"0{current_datetime.day}"
if current_datetime.month > 9:
    month = f"{current_datetime.month}"
else:
    month = f"0{current_datetime.month}"
year = f"{current_datetime.year}"
date = day + '.' + month + '.' + year

# decryption t2- with indexes(!)
if t2[5] != '' or t2[6] != '':
    for index in (7, 8, 9, 10, 11):
        t2[index] = t2[index].replace("\n", "")

# decryption t3
for t in range(last_hightonus_index, 1, -1):
    if t3[t] != '':
        t3[t] = t3[t].replace(",", ".")
        break

register = True
for t in range(1, hightonus_index):
    if t3[t] != '':
        register = False
if register:
    t3[hightonus_index] = t3[hightonus_index].capitalize()

# decryption t5
for t in range(len(t5)):
    if t5[t] == "Безусловные рефлексы: поисковый, сосательный, итд. ":
        t5[t] = (
                "\nБезусловные рефлексы: поисковый, сосательный, хоботковый, ладонно-рото-головной, хватательный, Моро\n"
                + "опоры, автоматической ходьбы, защитный, ползания, Галанта, Переса- симметричны. ")

# decryption t6
for t in range(len(t6)):
    if t6[t] == "Синдром мышечной дистонии: с кодом по МКБ. " or t6[
        t] == "Синдром мышечной дистонии: без кода по МКБ. ":
        t6[t] = "Синдром мышечной дистонии. "

# decryption t8
for t in range(len(t8)):
    if t8[t] == "консультация в Медико-генетическом научном центре\n":
        t8[
            t] = "консультация в Медико-генетическом научном центре им. академика Н.П. Бочкова (г. Москва, ул. Москворечье, 1)\n"
    if t8[t] == "консультации в НИИ педиатрии\n":
        t8[t] = (
                "консультация в НИИ педиатрии им. академика Ю.Е. Вельтищева ФГАОУ ВО РНИМУ им. Н.И. Пирогова\n"
                + "        Минздрава России (г. Москва, ул. Талдомская, 2)\n")

# decryption t9
for t in range(len(t9)):
    if t9[t] == "чередовать ванны: крапива, ромашка, череда\n":
        t9[t] = "чередовать ванны: крапива, ромашка, череда- 2 раза в неделю до № 10-15\n"
    if t9[t] == "успокаивающие ванны\n":
        t9[
            t] = "успокаивающие ванны: хмель, или мелисса, или лаванда, или ромашка, или душица- до № 15, затем сменить траву\n"
    if t9[t] == "ванны с крапивой\n":
        t9[t] = (
                "ванны с крапивой: 2 раза в неделю до № 10-15\n"
                + "        заварить 2 ст. л. сушёной крапивы на маленькую ванну, 4 ст. л. на большую ванну\n")
    if t9[t] == "ванны с морской солью\n":
        t9[t] = (
                "ванны с морской солью: 2 раза в неделю до № 10-15\n"
                + "        5 грамм морской соли на литр воды (75 г на маленькую ванну), температура воды до 37 °С:\n"
                + "        ребёнка искупать, промакнуть, через 20 минут соль смыть с тела\n")
    if t9[t] == "солёно-хвойные ванны\n":
        t9[t] = (
                "солёно-хвойные ванны: 2 раза в неделю до № 10-15\n"
                + "        5 грамм морской соли на литр воды (75 г на маленькую ванну) и 2 капли хвойного масла,\n"
                + "        температура воды до 37 °С: ребёнка искупать, промакнуть, через 20 минут соль смыть с тела\n")

# decryption t10
paraffin = (
        "        (при температуре окружающей среды до 25 °С, при температуре тела ребёнка до 37 °С):\n"
        + "        парафин (2 части: 400 г) и озокерит (1 часть: 200 г) растопить на водяной бане (осторожно! парафин горит!),\n"
        + "        вылить в форму, выстеленную медицинской клеёнкой, слоем в 1.5-2.0 см толщиной, остудить до 40 °С,\n"
        + "        наложить на 20 минут, затем снять (при беспокойстве ребёнка- снять раньше) и надеть тёплые носочки, варежки\n")

for t in range(len(t10)):
    if t10[t] == "парафин на кисти № 5+5\n" and t10[t + 1] == '':
        t10[t] = t10[t] + paraffin
    if t10[t] == "парафиновые сапожки № 10\n":
        t10[t] = t10[t] + paraffin

if t10[0] != "":
    for last_word_index in range(common_index, 0, -1):
        if t10[last_word_index] != "":
            t10[last_word_index] = t10[last_word_index].replace(", ", " № 10\n")
            break

# decryption t12
for t in range(len(t12)):
    if t12[t] == "мочегонный чай":
        t12[t] = (
                "мочегонный чай: 1 месяц\n"
                + "        лист земляники лесной, лист брусники, лист чёрной смородины, ромашку, плоды шиповника смешать в равных\n"
                + "        долях, заварить 1 чайную ложку смеси в 100 мл кипятка, пить по 1 чайной ложке чая 3 раза в день\n")
    elif t12[t] == "лист брусники":
        t12[t] = (
                "лист брусники: 1 месяц\n"
                + "        заварить 1 чайную ложку листа брусники в 100 мл кипятка, настоять, пить по 3 мл 2 раза в день 1 месяц\n")
    elif t12[t] == "пантогам (Hopantenic acid) 100 мг/мл в питьевом растворе":
        t12[t] = (t12[
                      t] + ":\n        7 дней: _______________________\n        21 день: ______________________\n        7 дней: _______________________\n")
    elif t12[t] == "пантогам или пантокальцин (Hopantenic acid) 250 мг":
        t12[t] = (t12[
                      t] + ":\n        7 дней: _______________________\n        21 день: ______________________\n        7 дней: _______________________\n")
    else:
        t12[t] = (t12[t] + ":   " + "_" * (95 - len(t12[t])) + "\n")

# title
document.add_heading(
    "ИП Павлова Ольга Игоревна\nгор. Пенза, ул. Пушкина, 15\nтел.: (+7 8412) 99-44-40 и +7 962 399-95-89\n"
    + "e-mail: rastuzdorovim@yandex.ru\nИНН: 583606020397   ОГРНИП: 315583600003199\n\n", level=3)
p = document.add_paragraph()
p.add_run("Дата:   ").bold = True
p.add_run(date + "\n")
p.add_run("Пациент:   ").bold = True
p.add_run(name + "\n")
p.add_run("Дата рождения:   ").bold = True
p.add_run(birthday + "\n\n\n")
p.add_run("Жалобы:   ").bold = True
if petition == "":
    p.add_run("Активных жалоб нет.\n\n")
p.add_run(petition + "\n\n")
p.add_run("Получил лечение:   __________________________________________").bold = True

p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
if head != "" or chest != "":
    if head != "":
        p.add_run("ОГ: ").bold = True
        p.add_run(head + " см")
    if chest != "":
        p.add_run("   ОГР: ").bold = True
        p.add_run(chest + " см")
    if fonticulus != "" and fonticulus.isdigit():
        p.add_run("   БР: ").bold = True
        p.add_run(fonticulus + "x" + fonticulus + " мм, не напряжён")
    elif fonticulus != "":
        p.add_run("   БР: ").bold = True
        p.add_run(fonticulus)
if allergy == "":
    p.add_run("\nНаличие аллергических реакций отрицают")
p.add_run("\n" + allergy)

# final text from dictionary
p = document.add_paragraph()
p.add_run("Неврологический статус:   ").bold = True
p.add_run("В сознании. Общемозговой симптоматики нет. Рвоты, судорог при осмотре нет. ")
for text in t1:
    p.add_run(text)
p.add_run("\nЧН:   ").bold = True
for t in [t2, t3, t4, t5]:
    for text in t:
        p.add_run(text)
    p.add_run("\n")
p.add_run("\n\n")
p.add_run("Диагноз поставлен на основании жалоб, объективных данных и лабораторных обследований.\n\n")
if t6[0] == "Здоров. ":
    p.add_run("В неврологическом статусе: без очаговых знаков.\n")
p.add_run("Диагноз:   ").bold = True
for text in t6:
    p.add_run(text)
p.add_run("\nДиагноз по МКБ-10:   ").bold = True
for text in t7:
    p.add_run(text)

document.add_page_break()
p = document.add_paragraph()
p.add_run("\nРекомендовано обследование:   \n").bold = True
for text in t8:
    p.add_run(text)
p.add_run("\n")
p.add_run("Рекомендовано лечение:   \n").bold = True
for t in [t9, t12, t10]:
    for text in t:
        p.add_run(text)
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Врач-невролог:   ").bold = True
p.add_run("Павлова Ольга Игоревна")
p = document.add_paragraph()
p.add_run("\n\nЯвка на повторный осмотр:   ____________________________\n\nЭпидемиологический анамнез:   ").bold = True
p.add_run(
    "Нахождение в очаге инфекции мама отрицает. Наличие контактов с больными\n"
    + "ОРВИ с инфекцией, вызванной новым коронавирусом SARS-CoV2, за последние 14 дней отрицает. Наличие тесных\n"
    + "контактов с лицами, находящимися под наблюдением по инфекции, вызванной новым коронавирусом SARS-CoV2,\n"
    + "которые в последующем заболели, отрицает. Наличие тесных контактов с лицами, у которых лабораторно\n"
    + "подтверждён  COVID-19, отрицает. За пределы Пензенской области и РФ в течение последнего месяца не выезжали.\n"
    + "В контакте с лицами, вернувшимися из других регионов и стран, не были.\n"
    + "TBC, ВИЧ, сифилис в семье и у родственников отрицает.\n"
    + "Дисфункций кишечника за последние 3 недели не было.\n\n"
    + "Выписка выдана на руки.\n"
    + "С планом обследования и лечения согласен (-на), о последствиях предупреждён (-на):   __________________________")
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("пациент (законный представитель)\n")
p.add_run(date)

# save and open file
sname = name.split()
savename = f"{sname[0]} {sname[1]} {date[:6]}{date[8:]}"
document.save(f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Архив\\{savename}.docx")
os.startfile(f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Архив\\{savename}.docx")

# print card
if add_card:
    f = open("card_number.txt", "r+")
    card_number = f.read()
    f.seek(0)
    f.write(f"{int(card_number) + 1}")
    f.close()

    card = Document("outpatient_card.docx")
    style = card.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    par = card.paragraphs
    par[15].add_run(card_number).bold = True
    par[19].add_run(date)
    par[20].add_run(name)
    par[21].add_run(sex)
    par[22].add_run(birthday)
    par[23].add_run(registration)
    card.save(
        f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Амбулаторные карты\\Карта № {card_number} {sname[0]} {sname[1]}.docx")
    os.startfile(
        f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Амбулаторные карты\\Карта № {card_number} {sname[0]} {sname[1]}.docx")

if __name__ == "__main__":
    main()
